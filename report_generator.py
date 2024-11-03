from dataclasses import dataclass
from typing import Dict, List, Optional, Union, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.shared import Cm
from pathlib import Path
import logging

@dataclass
class Style:
    size: int
    bold: bool
    color: RGBColor
    alignment: Optional[WD_ALIGN_PARAGRAPH] = None

class DocumentStyle:
    def __init__(self):
        self.styles = {
            'title': Style(size=16, bold=True, color=RGBColor(0, 0, 0),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER),
            'heading1': Style(size=14, bold=True, color=RGBColor(31, 73, 125)),
            'heading2': Style(size=12, bold=True, color=RGBColor(31, 73, 125)),
            'normal': Style(size=11, bold=False, color=RGBColor(0, 0, 0))
        }

class ReportGenerator:
    def __init__(self, custom_styles: Optional[Dict[str, Style]] = None):
        self.document = Document()
        self.style_manager = DocumentStyle()
        if custom_styles:
            self.style_manager.styles.update(custom_styles)
        
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)

    def add_full_page_image(self, image_path: Union[str, Path]) -> bool:
        """
        Adiciona uma imagem cobrindo toda a primeira página do documento.
        
        Args:
            image_path: Caminho para a imagem
            
        Returns:
            bool: True se a imagem foi adicionada com sucesso
        """
        try:
            image_path = Path(image_path)
            if not image_path.exists():
                raise FileNotFoundError(f"Imagem não encontrada: {image_path}")

            # Adicionar um parágrafo vazio para a imagem
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()

            # Configurar o tamanho da página
            section = self.document.sections[0]
            
            # Remover todas as margens da seção
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            
            # Configurar o parágrafo para remover recuos e espaçamentos
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0)
            paragraph_format.right_indent = Inches(0)
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = 1.0

            # Adicionar a imagem com o tamanho da página
            width = section.page_width
            height = section.page_height
            run.add_picture(str(image_path), width=width, height=height)

            # Restaurar as margens padrão para a próxima seção após a quebra de página
            self.document.add_section()
            new_section = self.document.sections[-1]
            new_section.left_margin = Inches(1)
            new_section.right_margin = Inches(1)
            new_section.top_margin = Inches(1)
            new_section.bottom_margin = Inches(1)

            return True
        
        except Exception as e:
            self.logger.error(f"Erro ao adicionar imagem de página inteira: {str(e)}")
            return False

    def apply_style(self, paragraph: Any, style_name: str) -> None:
        try:
            style = self.style_manager.styles[style_name]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            font = run.font
            font.size = Pt(style.size)
            font.bold = style.bold
            font.color.rgb = style.color
            
            if style.alignment:
                paragraph.alignment = style.alignment
                
        except KeyError:
            self.logger.error(f"Estilo não encontrado: {style_name}")
        except Exception as e:
            self.logger.error(f"Erro ao aplicar estilo: {str(e)}")

    def add_section_content(self, section: Dict[str, Any], level: int = 0) -> None:
        if 'title' in section:
            style_name = 'heading1' if level == 0 else 'heading2'
            paragraph = self.document.add_paragraph(section['title'])
            self.apply_style(paragraph, style_name)

        if 'text' in section:
            paragraph = self.document.add_paragraph(section['text'])
            self.apply_style(paragraph, 'normal')

        if 'table' in section:
            table = self.document.add_table(rows=len(section['table']),
                                            cols=len(section['table'][0]))
            table.style = 'Table Grid'
            
            for i, row in enumerate(section['table']):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)

        for subsection in section.get('subsections', []):
            self.add_section_content(subsection, level + 1)

    def generate_report(self, template: Dict[str, Any], output_path: Union[str, Path],
                        cover_image_path: Optional[Union[str, Path]] = None) -> bool:
        """
        Gera o relatório baseado no template fornecido.
        
        Args:
            template: Dicionário com o template do relatório
            output_path: Caminho para salvar o arquivo
            cover_image_path: Caminho opcional para a imagem de capa
            
        Returns:
            bool: True se o relatório foi gerado com sucesso
        """
        try:
            # Adicionar imagem de capa se fornecida
            if cover_image_path:
                if not self.add_full_page_image(cover_image_path):
                    return False

            # Adicionar título
            if 'title' in template:
                paragraph = self.document.add_paragraph(template['title'])
                self.apply_style(paragraph, 'title')

            # Adicionar seções
            for section in template.get('sections', []):
                self.add_section_content(section)
            
            self.document.save(str(output_path))
            
            self.logger.info(f"Relatório gerado com sucesso: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Erro ao gerar relatório: {str(e)}")
            return False

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    template = {
        "title": "Relatório de Análise de Projeto",
        "sections": [
            {
                "title": "Resumo Executivo",
                "text": "Este relatório apresenta uma análise detalhada do projeto X...",
                "table": [
                    ["Métrica", "Valor", "Status"],
                    ["Prazo", "60 dias", "Em dia"],
                    ["Orçamento", "R$ 100.000", "Dentro do previsto"]
                ],
                "subsections": [
                    {
                        "title": "Objetivos",
                        "text": "Os principais objetivos deste projeto são..."
                    }
                ]
            }
        ]
    }

    generator = ReportGenerator()
    success = generator.generate_report(
        template=template,
        output_path="relatorio.docx",
        cover_image_path="cover_images/cover_page_1.jpg"
    )

    if not success:
        logging.error("Falha ao gerar relatório")