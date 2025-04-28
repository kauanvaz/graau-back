from typing import Union, Optional
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
from pathlib import Path
import logging
import zipfile
import os
import shutil
import tempfile

class ReportGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.INFO)
        
    def _insert_headings_recursively(self, doc, headings: list, index: int, level: int=1):
        """
        Insere os títulos e subtítulos no documento de forma recursiva.
        Args:
            headings: Lista de dicionários com os títulos e subtítulos.
            index: Índice onde o título será inserido.
            level: Nível do título (1 para Heading 1, 2 para Heading 2, etc.).
            
        Returns:
            None
        """
        # Se for uma lista vazia não faz nada
        if not headings:
            return index
        
        current_index = index
        for sec in headings:
            # Adiciona uma quebra de página antes de cada título de nível 1
            if level == 1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            
            doc.paragraphs.insert(index, doc.add_paragraph(sec["title"], style=f"Heading {level}"))
            
            current_index += 1
            # Chama recursivamente para os subtítulos
            self._insert_headings_recursively(doc=doc, headings=sec["subtitles"], index=current_index, level=level+1)
            
        return current_index
    
    def _get_signing_area_name(self, headings: list) -> str:
        list_headings_level_1 = [h["title"].lower() for h in headings]
        
        if "proposta de encaminhamentos" in list_headings_level_1:
            return "proposta de encaminhamentos"
        elif "conclusão" in list_headings_level_1:
            return "conclusão"
        
        return None
    
    def _add_content(self, doc, text=None, bold=False, color=None, alignment=None, font='Segoe UI', space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        
        if not text:
            return
        
        run = p.add_run(text)
        
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
        if alignment:
            p.alignment = alignment
        
        run.font.name = font

    def _add_signing_content(self, doc):
        # Espaço em branco
        for _ in range(5):
            self._add_content(doc)

        # Instrução
        self._add_content(doc, "Instrução:", bold=True, font='Segoe UI Semibold', space_after=8)
        self._add_content(doc, "[informar auditores signatários]", color=RGBColor(191, 143, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # Supervisão
        self._add_content(doc, "Supervisão:", bold=True, font='Segoe UI Semibold', space_after=8)
        self._add_content(doc, "(assinado digitalmente)", color=RGBColor(128, 128, 128), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        self._add_content(doc, "[Nome]", color=RGBColor(191, 143, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        self._add_content(doc, "Auditor(a) de Controle Externo", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        self._add_content(doc, "Chefe da {{divisao_origem_ajustada_divisao}}", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # Visto
        self._add_content(doc, "Visto:", bold=True, font='Segoe UI Semibold', space_after=8)
        self._add_content(doc, "(assinado digitalmente)", color=RGBColor(128, 128, 128), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        self._add_content(doc, "[Nome]", color=RGBColor(191, 143, 0), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
        self._add_content(doc, "Diretor(a) da {{divisao_origem_ajustada_diretoria}}", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # Parágrafo em branco
        self._add_content(doc)

    def generate_headings_from_structure(self, doc, headings: list):
        """
        Gera os tópicos a partir da estrutura fornecida.
        Os estilos já devem existir no template.
        Args:
            headings: Lista de dicionários com os títulos e subtítulos.
            
        Returns:
            None
        """
        
        # Localiza o marcador e substitui
        for i, paragraph in enumerate(doc.paragraphs):
            if "<CONTEUDO>" in paragraph.text:
                # Remove o marcador
                p = paragraph.clear()
                
                if not headings:
                    # Se não houver títulos, remove o marcador e sai
                    return
                
                # Começa a inserir o primeiro título a partir do paragrafo do marcador
                p.text = headings[0]["title"]
                p.style = "Heading 1"
                next_index = i + 1
                
                # Insere os subtítulos do primeiro título
                next_index = self._insert_headings_recursively(doc=doc, headings=headings[0]["subtitles"], index=next_index, level=2)

                # Insere os subtítulos restantes, se houver
                if len(headings) > 1:
                    self._insert_headings_recursively(doc=doc, headings=headings[1:], index=next_index, level=1)
                
                break
            
        assinaturas_area = self._get_signing_area_name(headings)
        
        if not assinaturas_area:
            return
        
        for i, paragraph in enumerate(doc.paragraphs):
            if assinaturas_area in paragraph.text.lower():
                # Adiciona a conteúdo da área de assinatura logo após o parágrafo
                doc.paragraphs.insert(i+1, self._add_signing_content(doc))
                
                break
            
    def replace_existing_image(self, docx_path: str, target_image_filename: str, new_image_path: Union[str, Path]) -> bool:
        """
        Replace an existing image in the DOCX file by modifying its underlying ZIP archive.
        Utiliza um diretório temporário que é automaticamente removido.
        
        Args:
            docx_path: Path to the DOCX file.
            target_image_filename: The filename of the image to be replaced (e.g., "image1.png").
            new_image_path: Path to the new image file.
            
        Returns:
            bool: True if the image was replaced successfully.
        """
        try:
            # Cria um diretório temporário que será automaticamente excluído
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extrai o conteúdo do DOCX para o diretório temporário
                with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                target_image_path = os.path.join(temp_dir, 'word', 'media', target_image_filename)
                if not os.path.exists(target_image_path):
                    raise FileNotFoundError(f"Target image '{target_image_filename}' not found in the DOCX file.")
                
                # Substitui a imagem antiga pela nova
                shutil.copy(str(new_image_path), target_image_path)
                
                # Reempacota o DOCX com os conteúdos modificados (sobrescrevendo o arquivo original)
                with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as new_docx:
                    for foldername, _, filenames in os.walk(temp_dir):
                        for filename in filenames:
                            file_path = os.path.join(foldername, filename)
                            arcname = os.path.relpath(file_path, temp_dir)
                            new_docx.write(file_path, arcname)
                
                self.logger.info(f"Image replaced successfully in {docx_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error replacing image: {str(e)}")
            return False

    def generate_report(self,
                        context: dict,
                        output_path: str,
                        cover_image_path: Optional[Union[str, Path]] = None,
                        target_image_filename: Optional[str] = "image1.png") -> bool:
        """
        Generates the report using the template and provided context.

        Args:
            context: Dictionary with the template context.
            output_path: Path to save the output file.
            cover_image_path: Optional path to the cover image.
            replace_image: If True, replaces an existing image in the DOCX file.
            target_image_filename: The filename of the image in the DOCX to replace (required if replace_image is True).
            
        Returns:
            bool: True if the report was generated successfully.
        """
        try:
            doc = DocxTemplate(self.template_path)

            self.generate_headings_from_structure(doc=doc.get_docx(), headings=context.get("seccoes", []))
            
            if cover_image_path:
                # Render the document and save it temporarily
                temp_output = "temp_report.docx"
                doc.render(context)
                doc.save(temp_output)

                if not target_image_filename:
                    raise ValueError("target_image_filename must be provided.")

                if not self.replace_existing_image(temp_output, target_image_filename, cover_image_path):
                    return False

                # Move the temporary file to the final output path
                shutil.move(temp_output, output_path)
            else:
                doc.render(context)
                doc.save(output_path)

            self.logger.info(f"Report generated successfully: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error generating report: {str(e)}")
            return False

if __name__ == "__main__":
    sharepoint_data = [
        {
            "unidades_fiscalizadas": "P. M. CIDADE",
            "n_processo_eTCE": "TC/XXXXXX/20XX",
            "n_processo_eTCE_processo_tipo": "CONTAS-TOMADA DE CONTAS ESPECIAL",
            "exercicios": "20XX, 20YY",
            "VRF": "R$ 100.000,00"
        }
    ]
    
    generator = ReportGenerator("src/templates/Relatório Padrão - GRAAU.docx")
    
    success = generator.generate_report(
        context=sharepoint_data[0],
        output_path="src/reports/report_example.docx",
        cover_image_path="src/cover_images/cover_page_2.jpg",
    )
